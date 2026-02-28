"""
Service để trích xuất text từ PDF và Word documents
"""
from typing import Optional
import PyPDF2
from docx import Document
from pathlib import Path
import logging

logger = logging.getLogger(__name__)


class TextExtractor:
    """Extract text from PDF and DOCX files"""
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """
        Extract text from PDF using PyPDF2
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text as string
        """
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                logger.info(f"Extracting text from PDF: {num_pages} pages")
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text()
                    text += f"\n--- Trang {page_num} ---\n"
                    text += page_text + "\n"
                
            return text.strip()
            
        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            raise Exception(f"Không thể đọc file PDF: {str(e)}")
    
    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """
        Extract text from DOCX using python-docx
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text as string
        """
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extract tables
            table_data = []
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(" | ".join(row_data))
            
            # Combine all text
            text = "\n".join(paragraphs)
            
            if table_data:
                text += "\n\n--- Bảng dữ liệu ---\n"
                text += "\n".join(table_data)
            
            logger.info(f"Extracted {len(paragraphs)} paragraphs")
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            raise Exception(f"Không thể đọc file Word: {str(e)}")
    
    @classmethod
    def extract(cls, file_path: str) -> Optional[str]:
        """
        Auto-detect file type and extract text
        
        Args:
            file_path: Path to file
            
        Returns:
            Extracted text or None if unsupported
        """
        path = Path(file_path)
        extension = path.suffix.lower()
        
        logger.info(f"Extracting text from: {path.name}")
        
        if extension == '.pdf':
            return cls.extract_from_pdf(str(path))
        elif extension in ['.docx', '.doc']:
            return cls.extract_from_docx(str(path))
        else:
            raise ValueError(f"Định dạng không hỗ trợ: {extension}")